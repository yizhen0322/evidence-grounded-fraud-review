"""Build the Sunway-formatted CP2 final report as a Word document."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BLACK = "000000"
HEADER_GREY = "D9D9D9"
PALE_GREY = "F2F2F2"
DARK_GREY = RGBColor(64, 64, 64)
CONTENT_WIDTH_DXA = 9026
TABLE_INDENT_DXA = 120


def set_style_font(style, name: str, size: float | None = None) -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        style.font.size = Pt(size)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=100, bottom=100, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char_1, instr_text, fld_char_2))


def set_page_number_format(section, number_format: str, start: int = 1) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    node = OxmlElement("w:pgNumType")
    node.set(qn("w:fmt"), number_format)
    node.set(qn("w:start"), str(start))
    sect_pr.append(node)


def configure_section(section, *, numbered: bool = True) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    for part in (section.header, section.footer):
        for paragraph in part.paragraphs:
            for child in list(paragraph._p):
                if child.tag != qn("w:pPr"):
                    paragraph._p.remove(child)
    if numbered:
        add_page_number(section.footer.paragraphs[0])


def configure_document(doc: Document) -> None:
    # Resolved design system: narrative_proposal with a named Sunway academic
    # override (A4, Times New Roman 12 pt, black headings, 1.5 line spacing).
    # The override is intentionally applied through styles rather than ad-hoc
    # formatting so the report remains consistent and easy to revise in Word.
    section = doc.sections[0]
    configure_section(section, numbered=False)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, "Times New Roman", 12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)

    for name, size, before, after in (
        ("Title", 18, 0, 12),
        ("Subtitle", 14, 0, 12),
        ("Heading 1", 16, 12, 12),
        ("Heading 2", 14, 12, 6),
        ("Heading 3", 12, 10, 4),
    ):
        style = styles[name]
        set_style_font(style, "Times New Roman", size)
        style.font.color.rgb = RGBColor.from_string(BLACK)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.0

    styles["Heading 1"].paragraph_format.page_break_before = True

    for name in ("Header", "Footer"):
        style = styles[name]
        set_style_font(style, "Times New Roman", 10)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        set_style_font(style, "Times New Roman", 12)
        style.paragraph_format.left_indent = Cm(0.75)
        style.paragraph_format.first_line_indent = Cm(-0.45)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    if "Figure Caption" not in styles:
        fig_style = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        fig_style = styles["Figure Caption"]
    set_style_font(fig_style, "Times New Roman", 10)
    fig_style.font.italic = True
    fig_style.font.color.rgb = DARK_GREY
    fig_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_style.paragraph_format.space_before = Pt(3)
    fig_style.paragraph_format.space_after = Pt(9)

    if "Table Caption" not in styles:
        table_caption = styles.add_style("Table Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_caption = styles["Table Caption"]
    set_style_font(table_caption, "Times New Roman", 10)
    table_caption.font.bold = True
    table_caption.font.color.rgb = RGBColor.from_string(BLACK)
    table_caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table_caption.paragraph_format.space_before = Pt(6)
    table_caption.paragraph_format.space_after = Pt(3)
    table_caption.paragraph_format.keep_with_next = True

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    set_style_font(callout, "Times New Roman", 11)
    callout.font.color.rgb = RGBColor.from_string(BLACK)
    callout.paragraph_format.left_indent = Cm(0.45)
    callout.paragraph_format.right_indent = Cm(0.25)
    callout.paragraph_format.space_before = Pt(5)
    callout.paragraph_format.space_after = Pt(8)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    set_style_font(code, "Courier New", 9)
    code.paragraph_format.left_indent = Cm(0.6)
    code.paragraph_format.right_indent = Cm(0.3)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(3)
    code.paragraph_format.line_spacing = 1.0


INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")


def add_inline(paragraph, text: str) -> None:
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def add_paragraph(doc: Document, text: str, style: str | None = None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    add_inline(p, text)
    return p


def add_front_list_entry(doc: Document, label: str, page: int | str) -> None:
    row = doc.add_paragraph()
    row.paragraph_format.space_before = Pt(0)
    row.paragraph_format.space_after = Pt(3)
    row.paragraph_format.line_spacing = 1.0
    row.paragraph_format.tab_stops.add_tab_stop(
        Cm(15.7),
        WD_TAB_ALIGNMENT.RIGHT,
        WD_TAB_LEADER.DOTS,
    )
    label_run = row.add_run(label)
    label_run.font.name = "Times New Roman"
    label_run.font.size = Pt(11)
    page_run = row.add_run(f"\t{page}")
    page_run.font.name = "Times New Roman"
    page_run.font.size = Pt(11)


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator_row(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    lengths = [
        max(4, max((len(row[col]) if col < len(row) else 0) for row in rows))
        for col in range(cols)
    ]
    total = sum(lengths)
    widths = [max(900, round(CONTENT_WIDTH_DXA * length / total)) for length in lengths]
    correction = CONTENT_WIDTH_DXA - sum(widths)
    widths[-1] += correction

    # Table 2.1 contains short category labels in its first two columns and
    # comparatively long explanatory prose in the remaining columns. Purely
    # character-proportional sizing made the first heading wrap inside words
    # (for example, "Literatu re"). Give those labels a readable minimum while
    # preserving the full content width.
    is_literature_table = rows[0][0].strip().lower() == "literature stream" and cols == 4
    if is_literature_table:
        widths = [1600, 1700, 2700, CONTENT_WIDTH_DXA - 6000]

    # Table 2.2 compares four compact delivery-boundary properties across five
    # studies. Reserve enough width for the final measurement column and keep
    # the study label readable instead of relying on character-proportional
    # sizing for repeated "Not reported" cells.
    is_prior_work_table = rows[0][0].strip().lower() == "study" and cols == 5
    if is_prior_work_table:
        widths = [1250, 1800, 1800, 1800, CONTENT_WIDTH_DXA - 6650]

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    set_repeat_table_header(table.rows[0])
    keep_entire_table = rows[0][0].strip().lower() in {"research question", "rq"}
    for row_index, row in enumerate(rows):
        prevent_row_split(table.rows[row_index])
        for col_index in range(cols):
            cell = table.cell(row_index, col_index)
            cell.width = Inches(widths[col_index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[col_index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if is_literature_table or is_prior_work_table:
                set_cell_margins(cell, top=55, start=70, bottom=55, end=70)
            else:
                set_cell_margins(cell)
            text = row[col_index] if col_index < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, text)
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(
                    7.5
                    if is_prior_work_table
                    else 8
                    if is_literature_table
                    else 9
                    if cols >= 4
                    else 10
                )
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
            if row_index == 0:
                set_cell_shading(cell, HEADER_GREY)
            elif row_index % 2 == 0:
                set_cell_shading(cell, PALE_GREY)
            if keep_entire_table and row_index < len(rows) - 1:
                set_keep_with_next(p)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


IMAGE_RE = re.compile(r"!\[(.+?)\]\((.+?)\)")


def add_image(doc: Document, repo_root: Path, line: str) -> None:
    match = IMAGE_RE.fullmatch(line.strip())
    if not match:
        return
    caption, relative = match.groups()
    image_path = (repo_root / relative).resolve()
    if not image_path.exists():
        add_paragraph(doc, f"[Missing figure: {image_path}]", style="Callout")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    width = Inches(6.05)
    # Keep the tall Investigation Workspace figure compact enough that the
    # closing methodology subsection does not become an orphan page before
    # Chapter 4. Other, wider workbench screenshots retain the full text width.
    if image_path.name == "workbench_investigation.png":
        width = Inches(5.20)
    run.add_picture(str(image_path), width=width)
    caption_p = add_paragraph(doc, caption, style="Figure Caption")
    set_keep_with_next(p)
    caption_p.paragraph_format.keep_together = True


def add_cover(doc: Document, lines: list[str]) -> int:
    title = lines[0][2:].strip()
    metadata: dict[str, str] = {}
    for line in lines[1:25]:
        cleaned = line.strip().replace("**", "")
        if ":" in cleaned:
            label, value = cleaned.split(":", 1)
            metadata[label.strip()] = value.strip()

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(30)

    for text, size in (("CAPSTONE PROJECT 2", 16), ("FINAL REPORT", 14)):
        p = add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER)
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            run.font.bold = True

    title_p = add_paragraph(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER)
    title_p.paragraph_format.space_before = Pt(26)
    title_p.paragraph_format.space_after = Pt(12)
    for run in title_p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(18)
        run.font.bold = True

    study_focus = metadata.get("Study focus")
    if study_focus:
        focus_p = add_paragraph(doc, study_focus, align=WD_ALIGN_PARAGRAPH.CENTER)
        focus_p.paragraph_format.space_after = Pt(24)
        for run in focus_p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.italic = True

    cover_lines = [
        "by",
        metadata.get("Student", "NG YI ZHEN"),
        metadata.get("Student ID", "23076003"),
        "",
        metadata.get("Programme", "Bachelor of Computer Science (Hons)"),
        "",
        f"Supervisor: {metadata.get('Supervisor', 'Dr Tang Tiong Yew')}",
        f"Semester: {metadata.get('Semester', 'April 2026')}",
        f"Academic Year: {metadata.get('Academic year', '2025/2026')}",
        f"Submission Date: {metadata.get('Submission date', '15 July 2026')}",
    ]
    for text in cover_lines:
        p = add_paragraph(doc, text or " ", align=WD_ALIGN_PARAGRAPH.CENTER)
        p.paragraph_format.space_after = Pt(4 if text else 8)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            if text in {metadata.get("Student"), metadata.get("Student ID")}:
                run.font.bold = True

    footer_lines = [
        metadata.get("School", "School of Computing and Artificial Intelligence"),
        metadata.get("Faculty", "Faculty of Engineering and Technology"),
        "Sunway University",
    ]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    for text in footer_lines:
        line_p = add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER)
        line_p.paragraph_format.space_after = Pt(2)
        for run in line_p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    index = 1
    while index < len(lines) and lines[index].strip() != "# Abstract":
        index += 1

    front = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(front, numbered=True)
    set_page_number_format(front, "lowerRoman", 1)
    return index


def page_number_for(
    page_map: dict[str, int],
    key: str,
    *,
    allow_placeholder_page_numbers: bool,
) -> int | str:
    if key in page_map:
        return page_map[key]
    if allow_placeholder_page_numbers:
        return "?"
    raise KeyError(
        f"Missing page-map entry for {key!r}. Build a pass-1 document with "
        "--allow-placeholder-page-numbers, extract its page map, then rebuild."
    )


def add_contents(
    doc: Document,
    headings: list[tuple[int, str]],
    page_map: dict[str, int],
    *,
    allow_placeholder_page_numbers: bool,
) -> None:
    p = doc.add_heading("Table of Contents", level=1)
    p.paragraph_format.page_break_before = True
    for level, title in headings:
        row = doc.add_paragraph()
        row.paragraph_format.left_indent = Cm(0.55 * (level - 1))
        row.paragraph_format.space_before = Pt(0)
        row.paragraph_format.space_after = Pt(2)
        row.paragraph_format.line_spacing = 1.0
        row.paragraph_format.tab_stops.add_tab_stop(
            Cm(15.7),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )
        title_run = row.add_run(title)
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(10.5)
        title_run.bold = level == 1
        page_run = row.add_run(
            f"\t{page_number_for(page_map, title, allow_placeholder_page_numbers=allow_placeholder_page_numbers)}"
        )
        page_run.font.name = "Times New Roman"
        page_run.font.size = Pt(10.5)


def enable_field_updates(doc: Document) -> None:
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def build(
    markdown_path: Path,
    output_path: Path,
    page_map_path: Path | None = None,
    *,
    allow_placeholder_page_numbers: bool = False,
) -> None:
    repo_root = markdown_path.parents[2]
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    page_map = (
        json.loads(page_map_path.read_text(encoding="utf-8"))
        if page_map_path is not None and page_map_path.exists()
        else {}
    )
    if not page_map and not allow_placeholder_page_numbers:
        raise ValueError(
            "A page map is required for a submission build. Use "
            "--allow-placeholder-page-numbers only for the first render used "
            "to extract page numbers."
        )
    mainmatter_seen = False
    toc_headings: list[tuple[int, str]] = []
    for source_line in lines:
        source = source_line.strip()
        if source == "<!-- mainmatter -->":
            mainmatter_seen = True
            continue
        if not mainmatter_seen:
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", source)
        if heading:
            toc_headings.append((len(heading.group(1)), heading.group(2).strip()))
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = (
        "Evidence-Grounded Local-LLM Explanations for Credit Card Fraud Alert Review"
    )
    doc.core_properties.author = "NG YI ZHEN"
    doc.core_properties.subject = "Sunway University Capstone Project 2 Final Report"
    index = add_cover(doc, lines)

    suppress_next_h1_pagebreak = True
    manual_break_before_next_h1 = False
    current_h1 = ""
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped == "<!-- pagebreak -->":
            page_break = doc.add_paragraph()
            page_break.add_run().add_break(WD_BREAK.PAGE)
            # Heading 1 already has page_break_before in the document style.
            # Suppress that automatic break when a manual break immediately
            # precedes it, otherwise Word can render an unintended blank page.
            manual_break_before_next_h1 = True
            index += 1
            continue
        if stripped == "<!-- mainmatter -->":
            add_contents(
                doc,
                toc_headings,
                page_map,
                allow_placeholder_page_numbers=allow_placeholder_page_numbers,
            )
            main = doc.add_section(WD_SECTION.NEW_PAGE)
            configure_section(main, numbered=True)
            set_page_number_format(main, "decimal", 1)
            suppress_next_h1_pagebreak = True
            manual_break_before_next_h1 = False
            index += 1
            continue
        if stripped.startswith("# "):
            current_h1 = stripped[2:].strip()
            p = doc.add_heading(current_h1, level=1)
            if suppress_next_h1_pagebreak or manual_break_before_next_h1:
                p.paragraph_format.page_break_before = False
                suppress_next_h1_pagebreak = False
                manual_break_before_next_h1 = False
            index += 1
            continue
        # A manual break only replaces Heading 1's automatic break when the
        # heading is the next actual content item. Tables, paragraphs, or
        # lower-level headings consume the manual break normally.
        manual_break_before_next_h1 = False
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            index += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            index += 1
            continue
        if stripped.startswith("!["):
            add_image(doc, repo_root, stripped)
            index += 1
            continue
        if stripped.startswith("```"):
            language = stripped[3:].strip()
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index].rstrip("\n"))
                index += 1
            if index < len(lines):
                index += 1
            if language:
                label = add_paragraph(doc, language.upper(), style="Code Block")
                label.runs[0].bold = True
            for code_line in code_lines:
                add_paragraph(doc, code_line or " ", style="Code Block")
            continue
        if current_h1 == "List of Tables" and stripped.startswith("Table "):
            add_front_list_entry(
                doc,
                stripped,
                page_number_for(
                    page_map,
                    stripped,
                    allow_placeholder_page_numbers=allow_placeholder_page_numbers,
                ),
            )
            index += 1
            continue
        if current_h1 == "List of Figures" and stripped.startswith("Figure "):
            add_front_list_entry(
                doc,
                stripped,
                page_number_for(
                    page_map,
                    stripped,
                    allow_placeholder_page_numbers=allow_placeholder_page_numbers,
                ),
            )
            index += 1
            continue
        if current_h1 != "List of Tables" and re.match(
            r"^Table (?:\d+(?:\.\d+)?|[A-Z]\.\d+)\.\s+", stripped
        ):
            add_paragraph(doc, stripped, style="Table Caption")
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and is_separator_row(lines[index + 1]):
            rows = [split_table_row(stripped)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(split_table_row(lines[index]))
                index += 1
            add_table(doc, rows)
            continue
        if stripped.startswith("> "):
            callout_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                callout_lines.append(lines[index].strip().lstrip("> "))
                index += 1
            p = add_paragraph(doc, " ".join(callout_lines), style="Callout")
            p_pr = p._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "14")
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), BLACK)
            borders.append(left)
            p_pr.append(borders)
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            p = add_paragraph(doc, f"{numbered.group(1)}. {numbered.group(2)}")
            p.paragraph_format.left_indent = Cm(0.65)
            p.paragraph_format.first_line_indent = Cm(-0.45)
            index += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            add_paragraph(doc, bullet.group(1), style="List Bullet")
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            nxt = lines[index].strip()
            if not nxt or nxt.startswith(("#", "|", ">", "![", "<!--")) or re.match(r"^(\d+)\.\s+", nxt) or re.match(r"^[-*]\s+", nxt):
                break
            paragraph_lines.append(nxt)
            index += 1
        paragraph_text = " ".join(paragraph_lines)
        p = add_paragraph(doc, paragraph_text)
        if current_h1 == "References":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.first_line_indent = Cm(-1.27)
            p.paragraph_format.space_after = Pt(8)
        elif re.search(r"`[0-9a-f]{40,}`", paragraph_text):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    output_path.parent.mkdir(parents=True, exist_ok=True)
    enable_field_updates(doc)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("markdown", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--page-map", type=Path)
    parser.add_argument("--allow-placeholder-page-numbers", action="store_true")
    args = parser.parse_args()
    build(
        args.markdown.resolve(),
        args.output.resolve(),
        args.page_map.resolve() if args.page_map else None,
        allow_placeholder_page_numbers=args.allow_placeholder_page_numbers,
    )


if __name__ == "__main__":
    main()

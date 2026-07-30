"""Build the CP2 logbook in the official Sunway weekly-entry format."""

from __future__ import annotations

import argparse
import re
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table


BLACK = RGBColor(0, 0, 0)


@dataclass(frozen=True)
class WeekEntry:
    number: int
    title: str
    date_range: str
    objective: str
    activities: str
    evidence: str
    challenge: str
    reflection: str
    next_action: str


SOLUTION_ACTIONS = {
    1: (
        "I separated the work into a detector-performance question and an "
        "explanation-faithfulness question, with the detector as the prediction "
        "source, SHAP as evidence, and the LLM only as an articulation layer."
    ),
    2: (
        "I mapped design, implementation, evaluation, writing, and demonstration "
        "tasks to the official CP2 weeks, while leaving the title open to "
        "evidence-driven revision."
    ),
    3: (
        "I reframed the contribution around local generation, minimised evidence, "
        "deterministic validation, measured failure rates, and fail-closed fallback."
    ),
    4: (
        "I fixed the six experimental groups, five seeds, shared split, and "
        "multi-metric evaluation before training so the original hybrid hypothesis "
        "could fail fairly."
    ),
    5: (
        "I introduced content deduplication, stable case IDs, train-only "
        "preprocessing and resampling, validation-only selection, and manifest-based "
        "checks."
    ),
    6: (
        "I placed every detector group behind the same XGBoost interface, split, "
        "seeds, metrics, and artifact contract instead of favouring the hybrid "
        "configuration."
    ),
    7: (
        "I separated raw-output measurement from delivered-output policy and rejected "
        "repair or normalisation of failed LLM text."
    ),
    8: (
        "I organised the application around one analyst workflow: select an alert, "
        "inspect evidence, review validation, and record an action."
    ),
    9: (
        "I defined run manifests, stable source-run links, frozen artifact paths, and "
        "targeted tests at leakage, provenance, validation, and workflow boundaries."
    ),
    10: (
        "I retained the negative Autoencoder result, froze G6 seed 42 for downstream "
        "explanation work, and used the first workbench review to identify "
        "communication gaps."
    ),
    11: (
        "I added the requested result bar charts, retained the frozen detector after "
        "exploratory search, and positioned S0 as a readable semantic context rather "
        "than a competing benchmark."
    ),
    12: (
        "I completed the human pilot analysis, 49-item manual audit, report, logbook, "
        "presentation, verification records, and submission package without "
        "overstating the evidence."
    ),
}


SUPERVISOR_FEEDBACK = {
    11: (
        "Show more bar charts, especially for F1, and clarify the purpose of the "
        "website and the value of Ollama."
    )
}


def collapse(text: str) -> str:
    text = text.replace("`", "").replace("**", "")
    return re.sub(r"\s+", " ", text).strip()


def field(section: str, label: str) -> str:
    pattern = rf"\*\*{re.escape(label)}:\*\*\s*(.*?)(?=\n+\*\*|\n## |\Z)"
    match = re.search(pattern, section, flags=re.DOTALL)
    if not match:
        raise ValueError(f"Missing field {label!r}")
    return collapse(match.group(1))


def field_any(section: str, *labels: str) -> str:
    for label in labels:
        try:
            return field(section, label)
        except ValueError:
            continue
    raise ValueError(f"Missing all alternative fields: {labels!r}")


def parse_weeks(source: Path) -> list[WeekEntry]:
    text = source.read_text(encoding="utf-8")
    matches = list(re.finditer(r"^## Week (\d+)\. (.+)$", text, flags=re.MULTILINE))
    if len(matches) != 12:
        raise ValueError(f"Expected 12 weekly entries, found {len(matches)}")

    entries: list[WeekEntry] = []
    for index, match in enumerate(matches):
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        section = text[start:end]
        entries.append(
            WeekEntry(
                number=int(match.group(1)),
                title=collapse(match.group(2)),
                date_range=field(section, "Date"),
                objective=field(section, "Weekly objective"),
                activities=field(section, "Progress and output"),
                evidence=field(section, "Evidence/files"),
                challenge=field(section, "Problem or decision"),
                reflection=field(section, "Critical reflection"),
                next_action=field_any(section, "Next-week action", "Next action"),
            )
        )
    return entries


def set_run_font(
    run,
    size: float,
    *,
    name: str = "Cambria",
    bold: bool = False,
    italic: bool = False,
    color=BLACK,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_style_monochrome(doc: Document, style_name: str, *, font_name: str) -> None:
    style = doc.styles[style_name]
    style.font.name = font_name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font_name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font_name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)
    style.font.color.rgb = BLACK

    p_pr = style._element.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is not None:
        for border in borders:
            border.set(qn("w:color"), "000000")
            border.attrib.pop(qn("w:themeColor"), None)


def format_cover(doc: Document) -> None:
    set_style_monochrome(doc, "Title", font_name="Arial")
    set_style_monochrome(doc, "Heading 1", font_name="Arial")

    for paragraph in doc.paragraphs[:6]:
        for run in paragraph.runs:
            current_size = run.font.size.pt if run.font.size else 12.0
            set_run_font(
                run,
                current_size,
                name="Arial",
                bold=bool(run.bold),
                italic=bool(run.italic),
            )

    cover = doc.tables[0]
    for row in cover.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    current_size = run.font.size.pt if run.font.size else 12.0
                    set_run_font(
                        run,
                        current_size,
                        name="Arial",
                        bold=bool(run.bold),
                        italic=bool(run.italic),
                    )


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_cell_text(cell, text: str, size: float = 11.5, *, name: str = "Cambria") -> None:
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size, name=name)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, *, top: int = 70, start: int = 90, bottom: int = 70, end: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_layout(table: Table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Inches(1.55), Inches(4.70))
    for row in table.rows:
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]
        for cell in row.cells:
            set_cell_margins(cell)
    for row in table.rows:
        set_cell_text(row.cells[0], row.cells[0].text.strip(), 11.0)


def add_compact_paragraph(doc: Document, label: str, text: str, *, size: float = 10.5, after: float = 4.0):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.keep_together = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    label_run = paragraph.add_run(label)
    set_run_font(label_run, size, bold=True)
    if text:
        text_run = paragraph.add_run(" " + text)
        set_run_font(text_run, size)
    return paragraph


def add_week_heading(doc: Document, entry: WeekEntry, *, include_template_heading: bool) -> None:
    if include_template_heading:
        heading = doc.add_paragraph(style="Heading 1")
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(8)
        heading_run = heading.add_run("Weekly Entry Template")
        set_run_font(heading_run, 14.0, bold=True)


def split_activity_bullets(text: str, max_bullets: int = 4) -> list[str]:
    sentences = [
        sentence.strip()
        for sentence in re.split(r"(?<=[.!?])\s+(?=[A-Z0-9])", text)
        if sentence.strip()
    ]
    if len(sentences) <= max_bullets:
        return sentences
    groups: list[list[str]] = [[] for _ in range(max_bullets)]
    for index, sentence in enumerate(sentences):
        groups[min(index * max_bullets // len(sentences), max_bullets - 1)].append(sentence)
    return [" ".join(group) for group in groups if group]


def add_activity_bullet(doc: Document, text: str, *, italic: bool = False, size: float = 10.2) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.28)
    paragraph.paragraph_format.first_line_indent = Inches(-0.16)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    set_run_font(run, size, italic=italic)


def add_signature_table(doc: Document, signature_path: Path | None) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Inches(1.72), Inches(1.12), Inches(1.58), Inches(1.88))
    for column, width in zip(table.columns, widths, strict=True):
        column.width = width

    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = borders.find(qn(f"w:{edge}"))
                if tag is None:
                    tag = OxmlElement(f"w:{edge}")
                    borders.append(tag)
                tag.set(qn("w:val"), "nil")

    labels = ("Supervisor’s Signature:", "________________", "Student’s Signature:")
    for cell, text in zip(table.rows[0].cells[:3], labels, strict=True):
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, 10.0)

    student = table.rows[0].cells[3].paragraphs[0]
    student.alignment = WD_ALIGN_PARAGRAPH.CENTER
    student.paragraph_format.space_after = Pt(0)
    if signature_path is None:
        line = student.add_run("____________________")
        set_run_font(line, 10.0)
    else:
        student.add_run().add_picture(str(signature_path), width=Inches(1.08))


def remove_template_body(doc: Document) -> None:
    body = doc._element.body
    children = list(body.iterchildren())
    # Retain the official cover through the cover bookmark end, and retain sectPr.
    for child in children[9:-1]:
        body.remove(child)


def add_week(
    doc: Document,
    entry: WeekEntry,
    metadata_table_xml,
    signature_path: Path | None,
    *,
    first: bool,
) -> None:
    if first:
        page_break = doc.add_paragraph()
        page_break.add_run().add_break(WD_BREAK.PAGE)
    else:
        doc.add_page_break()

    add_week_heading(doc, entry, include_template_heading=first)

    table_xml = deepcopy(metadata_table_xml)
    doc._element.body.insert(-1, table_xml)
    table = Table(table_xml, doc)
    set_repeat_table_layout(table)
    set_cell_text(table.rows[0].cells[1], str(entry.number), 11.0)
    set_cell_text(table.rows[1].cells[1], entry.date_range, 11.0)

    add_compact_paragraph(doc, "Objectives for the Week:", entry.objective)
    add_compact_paragraph(doc, "Activities Completed:", "", after=2.0)
    for bullet in split_activity_bullets(entry.activities):
        add_activity_bullet(doc, bullet)
    add_activity_bullet(doc, "Supporting evidence: " + entry.evidence, italic=True, size=9.2)

    add_compact_paragraph(doc, "Challenges Encountered (if any):", entry.challenge)
    add_compact_paragraph(doc, "Solutions / Actions Taken (if any):", SOLUTION_ACTIONS[entry.number])
    add_compact_paragraph(doc, "Progress Reflection:", entry.reflection)
    add_compact_paragraph(doc, "Plans for Next Week:", entry.next_action)
    add_compact_paragraph(
        doc,
        "Supervisor's Feedback (if any):",
        SUPERVISOR_FEEDBACK.get(entry.number, "________________"),
        after=2.0,
    )
    add_signature_table(doc, signature_path)


def build(source: Path, template: Path, output: Path, signature_path: Path | None) -> None:
    entries = parse_weeks(source)
    reference = Document(template)
    metadata_table_xml = deepcopy(reference.tables[1]._tbl)
    doc = Document(template)

    cover = doc.tables[0]
    cover_values = (
        "NG YI ZHEN",
        "23076003",
        "Bachelor of Computer Science (Hons)",
        "Evidence-Grounded Local-LLM Explanations for Credit Card Fraud Alert Review",
        "Dr Tang Tiong Yew",
    )
    for index, (row, value) in enumerate(zip(cover.rows, cover_values, strict=True)):
        set_cell_text(row.cells[1], value, 10.8 if index == 3 else 11.5, name="Arial")

    format_cover(doc)

    remove_template_body(doc)
    for index, entry in enumerate(entries):
        add_week(doc, entry, metadata_table_xml, signature_path, first=index == 0)

    doc.core_properties.title = "Capstone Project 2 Logbook"
    doc.core_properties.author = "NG YI ZHEN"
    doc.core_properties.subject = "Sunway University Capstone Project 2"
    doc.core_properties.comments = "Built from the official Sunway Capstone Project Logbook template."
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Built {output}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("template", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--signature", type=Path)
    args = parser.parse_args()
    build(
        args.source.resolve(),
        args.template.resolve(),
        args.output.resolve(),
        args.signature.resolve() if args.signature else None,
    )


if __name__ == "__main__":
    main()

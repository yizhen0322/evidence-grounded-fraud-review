"""Build the formal CP2 supervisor meeting draft from Markdown."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BLACK = "000000"
HEADER_GREY = "D9D9D9"
PALE_GREY = "F2F2F2"
INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
IMAGE_RE = re.compile(r"!\[(.+?)\]\((.+?)\)")


def set_page_number(section, *, start: int = 1, roman: bool = False) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))
    page_type = OxmlElement("w:pgNumType")
    page_type.set(qn("w:fmt"), "lowerRoman" if roman else "decimal")
    page_type.set(qn("w:start"), str(start))
    section._sectPr.append(page_type)


def configure_section(section, *, numbered: bool = True, roman: bool = False) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    if numbered:
        set_page_number(section, roman=roman)


def configure_document(doc: Document) -> None:
    configure_section(doc.sections[0], numbered=False)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    for name, size in (("Title", 18), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    doc.styles["Heading 1"].paragraph_format.page_break_before = True

    if "Figure Caption" not in doc.styles:
        caption = doc.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = doc.styles["Figure Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(10)
    caption.font.italic = True
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)

    if "Table Caption" not in doc.styles:
        table_caption = doc.styles.add_style("Table Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_caption = doc.styles["Table Caption"]
    table_caption.font.name = "Times New Roman"
    table_caption.font.size = Pt(10)
    table_caption.font.bold = True
    table_caption.paragraph_format.keep_with_next = True
    table_caption.paragraph_format.space_before = Pt(6)
    table_caption.paragraph_format.space_after = Pt(3)

    if "Callout" not in doc.styles:
        callout = doc.styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = doc.styles["Callout"]
    callout.font.name = "Times New Roman"
    callout.font.size = Pt(11)
    callout.font.italic = True
    callout.paragraph_format.left_indent = Cm(0.6)
    callout.paragraph_format.right_indent = Cm(0.4)
    callout.paragraph_format.space_after = Pt(8)


def add_inline(paragraph, text: str) -> None:
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def add_paragraph(doc: Document, text: str, *, style: str | None = None, centered: bool = False):
    paragraph = doc.add_paragraph(style=style)
    if centered:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(paragraph, text)
    return paragraph


def split_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator(line: str) -> bool:
    cells = split_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def shade_cell(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    properties.append(shading)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_properties.append(repeat)

    for row_index, row in enumerate(rows):
        row_properties = table.rows[row_index]._tr.get_or_add_trPr()
        no_split = OxmlElement("w:cantSplit")
        row_properties.append(no_split)
        for column_index in range(columns):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(0)
            value = row[column_index] if column_index < len(row) else ""
            add_inline(paragraph, value)
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(8.5 if columns >= 4 else 9.5)
                run.bold = row_index == 0
            if row_index == 0:
                shade_cell(cell, HEADER_GREY)
            elif row_index % 2 == 0:
                shade_cell(cell, PALE_GREY)
    doc.add_paragraph()


def add_image(doc: Document, source_dir: Path, line: str) -> None:
    match = IMAGE_RE.fullmatch(line)
    if not match:
        return
    caption, relative_path = match.groups()
    image_path = (source_dir / relative_path).resolve()
    if not image_path.exists():
        add_paragraph(doc, f"Missing figure: {image_path}", style="Callout")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    width = Inches(5.25 if image_path.name == "03_investigation_workspace.png" else 6.05)
    paragraph.add_run().add_picture(str(image_path), width=width)
    add_paragraph(doc, caption, style="Figure Caption")


def add_static_toc(doc: Document) -> None:
    doc.add_page_break()
    heading = doc.add_heading("Table of Contents", level=1)
    heading.paragraph_format.page_break_before = False
    entries = (
        ("1. Introduction", 1),
        ("2. Literature Review", 5),
        ("3. Methodology and System Design", 9),
        ("4. Verified Preliminary Results", 16),
        ("References", 19),
    )
    for title, page in entries:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Cm(15.7), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        title_run = paragraph.add_run(title)
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(10)
        title_run.bold = True
        page_run = paragraph.add_run(f"\t{page}")
        page_run.font.name = "Times New Roman"
        page_run.font.size = Pt(10)


def add_cover(doc: Document, lines: list[str]) -> int:
    title = lines[0][2:].strip()
    metadata: dict[str, str] = {}
    for line in lines[1:30]:
        cleaned = line.strip().replace("**", "")
        if ":" in cleaned:
            key, value = cleaned.split(":", 1)
            metadata[key.strip()] = value.strip()

    doc.add_paragraph().paragraph_format.space_after = Pt(30)
    for text, size in (("CAPSTONE PROJECT 2", 16), ("SUPERVISOR REVIEW DRAFT", 14)):
        paragraph = add_paragraph(doc, text, centered=True)
        paragraph.paragraph_format.space_after = Pt(8)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            run.bold = True

    title_paragraph = add_paragraph(doc, title, centered=True)
    title_paragraph.paragraph_format.space_before = Pt(26)
    title_paragraph.paragraph_format.space_after = Pt(28)
    for run in title_paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(18)
        run.bold = True

    cover_lines = [
        "by",
        metadata.get("Student", "NG YI ZHEN"),
        metadata.get("Student ID", "23076003"),
        "",
        metadata.get("Programme", "Bachelor of Computer Science (Hons)"),
        "",
        f"Supervisor: {metadata.get('Supervisor', 'Dr Claymond Lim Wei Xiang')}",
        f"Semester: {metadata.get('Semester', 'April 2026')}",
        f"Academic Year: {metadata.get('Academic year', '2025/2026')}",
        f"Document Date: {metadata.get('Document date', '18 July 2026')}",
    ]
    for value in cover_lines:
        paragraph = add_paragraph(doc, value or " ", centered=True)
        paragraph.paragraph_format.space_after = Pt(4 if value else 8)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            if value in {metadata.get("Student"), metadata.get("Student ID")}:
                run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    for value in (
        metadata.get("School", "School of Computing and Artificial Intelligence"),
        metadata.get("Faculty", "Faculty of Engineering and Technology"),
        "Sunway University",
    ):
        paragraph = add_paragraph(doc, value, centered=True)
        paragraph.paragraph_format.space_after = Pt(2)

    index = 1
    while index < len(lines) and lines[index].strip() != "# Document Status":
        index += 1
    front = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(front, numbered=True, roman=True)
    return index


def enable_field_updates(doc: Document) -> None:
    settings = doc.settings.element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def build(markdown_path: Path, output_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "CP2 Supervisor Review Draft - Meeting 01"
    doc.core_properties.author = "NG YI ZHEN"
    doc.core_properties.subject = "Sunway University Capstone Project 2 supervisor review"
    index = add_cover(doc, lines)
    current_h1 = ""
    first_front_heading = True
    suppress_first_main_break = False

    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue
        if stripped == "<!-- pagebreak -->":
            doc.add_page_break()
            index += 1
            continue
        if stripped == "<!-- mainmatter -->":
            add_static_toc(doc)
            main = doc.add_section(WD_SECTION.NEW_PAGE)
            configure_section(main, numbered=True, roman=False)
            suppress_first_main_break = True
            index += 1
            continue
        if stripped.startswith("# "):
            current_h1 = stripped[2:].strip()
            paragraph = doc.add_heading(current_h1, level=1)
            if first_front_heading or suppress_first_main_break:
                paragraph.paragraph_format.page_break_before = False
                first_front_heading = False
                suppress_first_main_break = False
            index += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            index += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            index += 1
            continue
        if stripped.startswith("!["):
            add_image(doc, markdown_path.parent, stripped)
            index += 1
            continue
        if current_h1 != "References" and re.match(r"^Table \d+\.\d+\.\s+", stripped):
            add_paragraph(doc, stripped, style="Table Caption")
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and is_separator(lines[index + 1]):
            rows = [split_row(stripped)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(split_row(lines[index]))
                index += 1
            add_table(doc, rows)
            continue
        if stripped.startswith("> "):
            add_paragraph(doc, stripped[2:], style="Callout")
            index += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            paragraph = add_paragraph(doc, f"{numbered.group(1)}. {numbered.group(2)}")
            paragraph.paragraph_format.left_indent = Cm(0.65)
            paragraph.paragraph_format.first_line_indent = Cm(-0.45)
            index += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            add_paragraph(doc, bullet.group(1), style="List Bullet")
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if not next_line or next_line.startswith(("#", "|", ">", "![", "<!--")):
                break
            if re.match(r"^(\d+)\.\s+", next_line) or re.match(r"^[-*]\s+", next_line):
                break
            paragraph_lines.append(next_line)
            index += 1
        paragraph = add_paragraph(doc, " ".join(paragraph_lines))
        if current_h1 == "References":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Cm(1.27)
            paragraph.paragraph_format.first_line_indent = Cm(-1.27)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(4)
            for run in paragraph.runs:
                run.font.size = Pt(10.5)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    enable_field_updates(doc)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("markdown", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.markdown.resolve(), args.output.resolve())


if __name__ == "__main__":
    main()

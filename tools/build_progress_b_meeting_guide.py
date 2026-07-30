"""Build the private Progress B supervisor-preparation guide from Markdown."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(35, 39, 47)
GREY = RGBColor(95, 99, 104)
LIGHT_BLUE = "DCE6F1"
LIGHT_GREY = "F2F2F2"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in (
        ("Title", 25, BLUE),
        ("Heading 1", 17, BLUE),
        ("Heading 2", 13.5, DARK),
        ("Heading 3", 11.5, BLUE),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    doc.styles["List Bullet"].font.name = "Aptos"
    doc.styles["List Number"].font.name = "Aptos"
    doc.styles["List Bullet"]._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "PingFang SC")
    doc.styles["List Number"]._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "PingFang SC")


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Progress B Supervisor Meeting Preparation")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Credit Card Fraud Detection using a Hybrid Autoencoder-XGBoost Model with Local LLM Explanations"
    )
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = DARK

    doc.add_paragraph()
    for text in (
        "NG YI ZHEN",
        "Student ID: 23076003",
        "Supervisor: Dr Claymond Lim Wei Xiang",
        "Prepared: 20 July 2026",
    ):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(text).font.size = Pt(12)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Private study and consultation guide - not part of the academic report")
    run.italic = True
    run.font.color.rgb = GREY

    doc.add_page_break()


INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline(paragraph, text: str) -> None:
    position = 0
    for match in INLINE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(9)
            run.font.color.rgb = BLUE
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def split_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator(line: str) -> bool:
    cells = split_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(rows):
        for col_index in range(width):
            cell = table.cell(row_index, col_index)
            text = row[col_index] if col_index < len(row) else ""
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_inline(paragraph, text)
            for run in paragraph.runs:
                run.font.size = Pt(8.5)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for run in paragraph.runs:
                    run.bold = True
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GREY)
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph()


def add_quote(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.8)
    paragraph.paragraph_format.right_indent = Cm(0.5)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(7)
    add_inline(paragraph, text)
    for run in paragraph.runs:
        run.italic = True
        run.font.color.rgb = BLUE


def add_markdown(doc: Document, path: Path, *, skip_title: bool = False) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    index = 0
    skipped_h1 = False
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("# "):
            if skip_title and not skipped_h1:
                skipped_h1 = True
            else:
                doc.add_heading(stripped[2:].strip(), level=1)
            index += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            index += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            index += 1
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
            index += 1
            continue
        if stripped == "---":
            index += 1
            continue
        if stripped.startswith("| and "):
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and is_separator(lines[index + 1].strip()):
            rows = [split_row(stripped)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(split_row(lines[index]))
                index += 1
            add_table(doc, rows)
            continue
        if stripped.startswith("> "):
            quote_lines = [stripped[2:]]
            index += 1
            while index < len(lines) and lines[index].strip().startswith("> "):
                quote_lines.append(lines[index].strip()[2:])
                index += 1
            add_quote(doc, " ".join(quote_lines))
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            add_inline(paragraph, numbered.group(2))
            index += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet.group(1))
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate or candidate.startswith(("#", "|", ">", "- ")):
                break
            if re.match(r"^\d+\.\s+", candidate) or candidate == "---":
                break
            paragraph_lines.append(candidate)
            index += 1
        paragraph = doc.add_paragraph()
        add_inline(paragraph, " ".join(paragraph_lines))


def build(start_here: Path, questions: Path, output: Path) -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    contents = doc.add_heading("How to use this guide", level=1)
    contents.paragraph_format.space_after = Pt(6)
    for text in (
        "Part I contains the meeting story, exact figures, technical explanations and live-demo order.",
        "Part II contains sixty likely supervisor questions with concise answers.",
        "Use the short answer first. Add technical detail only when the supervisor asks a follow-up question.",
        "The academic report remains the source to give the supervisor; this guide is private preparation material.",
    ):
        doc.add_paragraph(text, style="List Bullet")
    doc.add_page_break()

    doc.add_heading("Part I - Meeting briefing", level=1)
    add_markdown(doc, start_here, skip_title=True)
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("Part II - Supervisor question bank", level=1)
    add_markdown(doc, questions, skip_title=True)

    for index, section in enumerate(doc.sections):
        if index > 0:
            section.footer.is_linked_to_previous = False
        footer = section.footer.paragraphs[0]
        footer._p.clear_content()
        add_page_number(footer)

    core = doc.core_properties
    core.title = "Progress B Supervisor Meeting Preparation"
    core.subject = "Private FYP consultation and viva preparation guide"
    core.author = "NG YI ZHEN"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("start_here", type=Path)
    parser.add_argument("questions", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.start_here.resolve(), args.questions.resolve(), args.output.resolve())


if __name__ == "__main__":
    main()

"""Build a clean Word support document from the CP2 Markdown deliverables."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_font(style, name: str, size: float) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def clean(text: str) -> str:
    return text.replace("**", "").replace("`", "")


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    add_page_number(section.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    set_font(normal, "Times New Roman", 11.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size in (("Title", 20), ("Subtitle", 13), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        style = doc.styles[style_name]
        set_font(style, "Times New Roman", size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(17, 35, 58)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        set_font(style, "Times New Roman", 11.5)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(3)


def build(source: Path, output: Path, title: str, subtitle: str) -> None:
    doc = Document()
    configure(doc)
    doc.core_properties.title = title
    doc.core_properties.author = "NG YI ZHEN"
    doc.core_properties.subject = "Sunway University Capstone Project 2"

    cover = doc.add_paragraph(style="Title")
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.add_run(title)
    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(subtitle)
    for _ in range(4):
        doc.add_paragraph()
    identity = doc.add_paragraph()
    identity.alignment = WD_ALIGN_PARAGRAPH.CENTER
    identity.add_run("NG YI ZHEN\n23076003\nBachelor of Computer Science (Hons)\nSunway University").bold = True
    doc.add_page_break()

    lines = source.read_text(encoding="utf-8").splitlines()
    first_section = True
    index = 0
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("# "):
            if index == 0:
                index += 1
                continue
            doc.add_heading(clean(stripped[2:]), level=1)
            index += 1
            continue
        if stripped.startswith("## "):
            heading = clean(stripped[3:])
            if not first_section and (
                heading.startswith("Entry ")
                or heading.startswith("Week ")
                or heading.startswith("Slide ")
                or heading.startswith("Rubric Alignment")
            ):
                doc.add_page_break()
            first_section = False
            doc.add_heading(heading, level=2)
            index += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(clean(stripped[4:]), level=3)
            index += 1
            continue
        if re.match(r"^[-*]\s+", stripped):
            doc.add_paragraph(clean(re.sub(r"^[-*]\s+", "", stripped)), style="List Bullet")
            index += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            doc.add_paragraph(clean(re.sub(r"^\d+\.\s+", "", stripped)), style="List Number")
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            nxt = lines[index].strip()
            if not nxt or nxt.startswith("#") or nxt.startswith("**") or re.match(r"^[-*]\s+", nxt) or re.match(r"^\d+\.\s+", nxt):
                break
            paragraph_lines.append(nxt)
            index += 1
        paragraph = doc.add_paragraph()
        text = clean(" ".join(paragraph_lines))
        label_match = re.match(r"^([^:]{2,80}:)\s*(.*)$", text)
        if label_match:
            paragraph.add_run(label_match.group(1)).bold = True
            if label_match.group(2):
                paragraph.add_run(" " + label_match.group(2))
            label = label_match.group(1).lower()
            if label.startswith("evidence/files") or label.startswith("existing visual recommendation"):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            paragraph.add_run(text)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Built {output}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--title", required=True)
    parser.add_argument("--subtitle", required=True)
    args = parser.parse_args()
    build(args.source.resolve(), args.output.resolve(), args.title, args.subtitle)


if __name__ == "__main__":
    main()

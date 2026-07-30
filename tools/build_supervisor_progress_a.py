"""Build the CP2 Chapters 1-3 academic working draft."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Cm, Pt

from build_supervisor_meeting_packet import (
    add_paragraph,
    add_table,
    configure_document,
    configure_section,
    enable_field_updates,
    is_separator,
    split_row,
)


def add_cover(doc: Document, lines: list[str]) -> int:
    title = lines[0][2:].strip()
    metadata: dict[str, str] = {}
    for line in lines[1:30]:
        cleaned = line.strip().replace("**", "")
        if ":" in cleaned:
            key, value = cleaned.split(":", 1)
            metadata[key.strip()] = value.strip()

    doc.add_paragraph().paragraph_format.space_after = Pt(30)
    paragraph = add_paragraph(doc, "CAPSTONE PROJECT 2", centered=True)
    paragraph.paragraph_format.space_after = Pt(8)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)
        run.bold = True

    title_paragraph = add_paragraph(doc, title, centered=True)
    title_paragraph.paragraph_format.space_before = Pt(28)
    title_paragraph.paragraph_format.space_after = Pt(28)
    for run in title_paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(18)
        run.bold = True

    cover_lines = (
        "by",
        metadata.get("Student", "NG YI ZHEN"),
        metadata.get("Student ID", "23076003"),
        "",
        metadata.get("Programme", "Bachelor of Computer Science (Hons)"),
        "",
        f"Supervisor: {metadata.get('Supervisor', 'Dr Claymond Lim Wei Xiang')}",
        f"Semester: {metadata.get('Semester', 'April 2026')}",
        f"Academic Year: {metadata.get('Academic year', '2025/2026')}",
    )
    for value in cover_lines:
        paragraph = add_paragraph(doc, value or " ", centered=True)
        paragraph.paragraph_format.space_after = Pt(4 if value else 8)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            if value in {metadata.get("Student"), metadata.get("Student ID")}:
                run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    for value in (
        metadata.get("School", "School of Computing and Artificial Intelligence"),
        metadata.get("Faculty", "Faculty of Engineering and Technology"),
        "Sunway University",
    ):
        paragraph = add_paragraph(doc, value, centered=True)
        paragraph.paragraph_format.space_after = Pt(2)

    index = 1
    while index < len(lines) and lines[index].strip() != "<!-- mainmatter -->":
        index += 1
    front = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(front, numbered=True, roman=True)
    return index


def add_static_toc(doc: Document) -> None:
    heading = doc.add_heading("Table of Contents", level=1)
    heading.paragraph_format.page_break_before = False
    entries = (
        ("1. Introduction", 1),
        ("2. Literature Review", 4),
        ("3. Methodology and System Design", 6),
        ("References", 8),
    )
    for title, page in entries:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Cm(15.7), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        title_run = paragraph.add_run(title)
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(10)
        title_run.bold = True
        page_run = paragraph.add_run(f"\t{page}")
        page_run.font.name = "Times New Roman"
        page_run.font.size = Pt(10)


def build(markdown_path: Path, output_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "CP2 Working Draft - Chapters 1 to 3"
    doc.core_properties.author = "NG YI ZHEN"
    doc.core_properties.subject = "Sunway University Capstone Project 2 academic working draft"
    index = add_cover(doc, lines)
    current_h1 = ""
    suppress_first_main_break = False

    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue
        if stripped == "<!-- pagebreak -->":
            doc.add_page_break()
            index += 1
            continue
        if stripped == "<!-- mainmatter -->":
            add_static_toc(doc)
            main = doc.add_section(WD_SECTION.NEW_PAGE)
            configure_section(main, numbered=True, roman=False)
            suppress_first_main_break = True
            index += 1
            continue
        if stripped.startswith("# "):
            current_h1 = stripped[2:].strip()
            paragraph = doc.add_heading(current_h1, level=1)
            if suppress_first_main_break:
                paragraph.paragraph_format.page_break_before = False
                suppress_first_main_break = False
            index += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            index += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            index += 1
            continue
        if current_h1 != "References" and re.match(r"^Table \d+\.\d+\.\s+", stripped):
            add_paragraph(doc, stripped, style="Table Caption")
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and is_separator(lines[index + 1]):
            rows = [split_row(stripped)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(split_row(lines[index]))
                index += 1
            add_table(doc, rows)
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            paragraph = add_paragraph(doc, f"{numbered.group(1)}. {numbered.group(2)}")
            paragraph.paragraph_format.left_indent = Cm(0.65)
            paragraph.paragraph_format.first_line_indent = Cm(-0.45)
            index += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            add_paragraph(doc, bullet.group(1), style="List Bullet")
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if not next_line or next_line.startswith(("#", "|", "<!--")):
                break
            if re.match(r"^(\d+)\.\s+", next_line) or re.match(r"^[-*]\s+", next_line):
                break
            paragraph_lines.append(next_line)
            index += 1
        paragraph = add_paragraph(doc, " ".join(paragraph_lines))
        if current_h1 == "References":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Cm(1.27)
            paragraph.paragraph_format.first_line_indent = Cm(-1.27)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(5)
            for run in paragraph.runs:
                run.font.size = Pt(10.5)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    enable_field_updates(doc)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("markdown", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.markdown.resolve(), args.output.resolve())


if __name__ == "__main__":
    main()
